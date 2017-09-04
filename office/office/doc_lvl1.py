__author__ = "jzs"
__copyright__ = "@COPYLEFT 2017 ALL WRONGS RESERVED"

from docx import Document
from docx.shared import Pt
import os, time, argparse
import hashlib, zlib

def md5(fname):
    hash_md5 = hashlib.md5()
    with open(fname, "rb") as f:
        for chunk in iter(lambda: f.read(2**16), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def sha1(fname):
    sha1sum = hashlib.sha1()
    with open(fname, "rb") as f:
        for chunk in iter(lambda: f.read(2**16), b""):
            sha1sum.update(chunk)
    return sha1sum.hexdigest()

def crc32(fname):
    __digest = 0
    with open(fname, "rb") as f:
        for chunk in iter(lambda: f.read(2**16), b""):
            __digest = zlib.crc32(chunk, __digest) & 0xffffffff 
    return '{:08x}'.format(__digest)


parser = argparse.ArgumentParser(description='Scan source folder and generate level 1 doc for FSTEK.')
parser.add_argument('-i', '--input-dir', type=str, metavar='folder_path', required=True,
                    help='input folder path that points to the root of source code. default is current folder')
parser.add_argument('-o', '--output-file', default='document_lvl1.docx', type=str, metavar='file_name', 
                    help='output file name. default is `document_level1.docx`')
parser.add_argument('-s', '--suffixes', nargs='*', default='.cpp', 
                    help='file suffix(es) that you want to put into doc. e.g. `--suffixes .cpp .h`')
parser.add_argument('--max', default=-1, metavar='number_of_files', type=int,
                    help='max number of source files you want to put into docx file.')
args = parser.parse_args()

source_list = []
for root, dirs, files in os.walk(args.input_dir):
    for file in files:
        if any(file.endswith(suffix) for suffix in args.suffixes):
            source_list.append((root, file))
file_cnt = len(source_list)
if args.max > 0:
    file_cnt = args.max

row_per_file = 3
col_per_file = 6
#row_cnt = file_cnt * row_per_file + 1

document = Document()

table = document.add_table(rows=1, cols=col_per_file)
table.style = 'Table Grid'
table.style.font.name = 'Times New Roman'
table.style.font.size = Pt(10)

#write the header:
run = table.cell(0, 0).add_paragraph().add_run()
run.text = '№ пп'
run.font.bold = True
run = table.cell(0, 1).add_paragraph().add_run()
run.text = 'Имя файла'
run.font.bold = True
run = table.cell(0, 2).add_paragraph().add_run()
run.text = 'Дата создания'
run.font.bold = True
run = table.cell(0, 3).add_paragraph().add_run()
run.text = 'Длина, байт'
run.font.bold = True
run = table.cell(0, 4).add_paragraph().add_run()
run.text = 'КС (алгоритм ВКС)'
run.font.bold = True
run = table.cell(0, 5).add_paragraph().add_run()
run.text = 'Выполняемая функция'
run.font.bold = True

for i in range(file_cnt):
    path, name = source_list[i]
    fname = os.path.join(path, name)

    cells = table.add_row().cells
    cells[0].text = str(i+1)
    cells[1].text = name
    (mode, ino, dev, nlink, uid, gid, size, atime, mtime, ctime) = os.stat(fname)
    cells[2].text = time.strftime("%d.%m.%y %H-%M", time.gmtime(ctime))
    length = str(size)
    checksum = crc32(fname)
    cells[3].text = length
    cells[4].text = checksum
    cells[5].text = 'description'

    cells = table.add_row().cells
    cells[0].merge(cells[2])
    cells[0].text = 'итого: файлов - 1'
    cells[3].text = length
    cells[4].text = checksum

    cells = table.add_row().cells
    cells[0].merge(cells[-1])
    cells[0].text = "Каталог %s" % (path)

    if (i+1) % 10 == 0:
        print("%d / %d = %3.2f%%" % (i+1, file_cnt, (i+1)/file_cnt * 100.0))

document.save(args.output_file)
